import os
from docx import Document
import FileLoad


diretorysave = r'Registro Suspenso\Enviar'
def createRegistro(cpf, name, card, value):
    doc = Document('modelregistro.docx')
    FileLoad.createDiretory()
    for paragraph in doc.paragraphs:
        if 'CPFCLIENTE' in paragraph.text:
            paragraph.text = paragraph.text.replace('CPFCLIENTE', cpf)
        elif 'NAMECLIENTE' in paragraph.text:
            paragraph.text = paragraph.text.replace('NAMECLIENTE', name)
        elif 'CARDCLIENTE' in paragraph.text:
            paragraph.text = paragraph.text.replace('CARDCLIENTE', card)
        elif 'VALUEPROTESTO' in paragraph.text:
            paragraph.text = paragraph.text.replace('VALUEPROTESTO', f'R$ {value}')
    files = os.listdir(diretorysave)
    filename = f'{name} - {card}.docx'
    if filename not in files:
        doc.save(f'{diretorysave}\\{filename}')

def getRegistro(diretory):
    doc = Document(diretory)
    getreturn = {'CPF':None,
                 'CONTRATO':None,
                 'VALOR':None}
    for paragraph in doc.paragraphs:
        if ':' in paragraph.text and len(paragraph.text) < 50:
            vectorp = str(paragraph.text).split(':')
            find = vectorp[0]
            inf = vectorp[1]
            if 'CPF' in find:
                getreturn['CPF'] = inf
            elif 'Contrato' in find:
                getreturn['CONTRATO'] = inf
            elif 'Valor' in find:
                getreturn['VALOR'] = inf
    return getreturn




